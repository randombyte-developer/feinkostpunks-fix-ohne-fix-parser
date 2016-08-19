import urllib2
from bs4 import BeautifulSoup
import re
import os
import sys
from docx import Document
from docx.shared import Pt

class Recipe:
	def __init__(self, title, ingredients, texts):
  		self.title = title
		self.ingredients = ingredients
		self.texts = texts

def get_html(link):
	return urllib2.urlopen(link).read()

def get_recipe_links(overview_page_soup):
	return overview_page_soup.select("div.pf-content li > a")
	

def get_title(recipe_soup):
	return recipe_soup.find("h1").string

def get_ingredients(recipe_soup):
	for strong_text in recipe_soup.select("div.pf-content p > strong"): #the 'Zutaten' headline is in a strong-tag
		if (strong_text.string):
			if ("Zutaten" in strong_text.string):
				for sibling in strong_text.parent.find_next_siblings(limit = 3): #within the next 3 siblings definitely is the ingredient's p-tag where the 'list' is
					if (len(sibling.find_all("br")) > 3): #the ingredients aren't a list, they are br-tag separated texts -> there have to be many br-tags
						return list(sibling.stripped_strings) #strip away all those br-tags, return as list
			else:
				print("Skipped 'NoneType' while searching for beginning of ingredients!")
	print("Couldn't find any ingredients!")
	return ["Keine Zutaten"]
					
def get_text(recipe_soup):
	for strong_text in recipe_soup.select("div.pf-content p > strong"): #the 'Zubereitung' headline is in a strong-tag
		if (strong_text.string):
	                if ("Zubereitung" in strong_text.string):
				texts = []
                	        for sibling in strong_text.parent.find_next_siblings("p"):
					for strong_text_2 in sibling.find_all("strong"):
						if (strong_text_2.string):
							if ("Kosten" in strong_text_2.string or "Preis" in strong_text_2.string or "Geld" in strong_text_2.string):
								return texts
						else:
							print("Skipped 'NoneType' while searching for end of instructions!")
					texts += sibling.stripped_strings
				print("There's probably the cost comparasion in the instructions and couldn't filtered out!")
				return texts
		else:
			print("Skipped 'NoneType' while searching for beginning of instructions!")
	print("Couldn't find any making instructions!")
	return ["Keine Anweisungen"]
					

def get_recipe(link):
	s = BeautifulSoup(get_html(link), "lxml")
	title = get_title(s)
	print("==================================================")
	print("Processing '%s'" % title)
	return Recipe(title, get_ingredients(s), get_text(s))

output_dir = os.path.join(os.path.dirname(os.path.realpath(sys.argv[0])), "output")
def create_output_dir():
	if (not os.path.exists(output_dir) or not os.path.isdir(output_dir)):
		os.makedirs(output_dir)

def generate_documents(recipes):
	create_output_dir()
	for recipe in recipes:
		doc = Document()

		heading_r = doc.add_paragraph().add_run(recipe.title.string)
		heading_r.font.size = Pt(14)
		heading_r.font.bold = True
		heading_r.font.underline = True

		ingredients_p = doc.add_paragraph()
		ingredients_r = ingredients_p.add_run("Zutaten")
		ingredients_r.font.size = Pt(13)
		ingredients_r.font.bold = True
		for ingredient in recipe.ingredients:
			ingredient_li_p = doc.add_paragraph(style = "ListBullet")
			ingredient_li_p.add_run(ingredient).font.size = Pt(13)

		instructions_p = doc.add_paragraph()
		instructions_r = instructions_p.add_run("Zubereitung")
		instructions_r.font.size = Pt(13)
		instructions_r.font.bold = True
		for text in recipe.texts:
			text_r = doc.add_paragraph().add_run(text)
			text_r.font.size = Pt(13)

		doc.save(os.path.join(output_dir, recipe.title.string + ".docx"))

#Entry point
root_soup = BeautifulSoup(get_html("http://feinkostpunks.de/fix-ohne-fix/"), "lxml")
#recipes = map(lambda link_a_tag: get_recipe(link_a_tag["href"]), get_recipe_links(root_soup))
recipes = [get_recipe("http://feinkostpunks.de/fix-ohne-fix-nudel-schinken-gratin/")]
generate_documents(recipes)
print("Done: " + output_dir)
